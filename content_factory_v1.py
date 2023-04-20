import streamlit as st
import openai
import pandas as pd
import time
import re
import os
import zipfile
from io import BytesIO
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from html.parser import HTMLParser

def create_url_path(keyword):
    url_path = keyword.lower()
    url_path = re.sub(r"[^a-z0-9\s]+", "", url_path)  # Remove non-alphanumeric and non-space characters
    url_path = url_path.replace(" ", "-")  # Replace spaces with hyphens

    # Remove trailing hyphen if present
    if url_path[-1] == "-":
        url_path = url_path[:-1]

    return f"/{url_path}"


def create_full_path(domain, url_path):
    return f"https://{domain}{url_path}"


def generate_content(api_key, prompt, sections):
    openai.api_key = api_key

    system_message = (
        """

You are an AI language model. Your task is to follow the provided outline and ensure that the content is well-structured, SEO-friendly, and addresses the key points in each section.
Make sure to use clear, concise language and provide practical advice, examples, and tips where applicable.

""" )
    
    print(f"Generated prompt:\n{prompt}\n")

    messages = [
        {"role": "system", "content": system_message},
        {"role": "user", "content": prompt}
    ]

    completion = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        max_tokens=3200,
        messages=messages
    )

    response = completion.choices[0].message.content.strip()
    return response


def generate_related_links(df, current_topic):
    current_category = df.loc[df['topic'] == current_topic, 'category'].values[0]
    current_full_path = df.loc[df['topic'] == current_topic, 'full path'].values[0]
    related_links = df[df['category'] == current_category][['topic', 'full path']]
    
    # Filter out the self-link by comparing the full paths
    related_links = related_links[related_links['full path'] != current_full_path]

    return related_links.to_dict('records')


st.cache_data(show_spinner=False, ttl=86400)
def generate_article(api_key, topic, sections, related_links, definition_only=False):
    if definition_only:
        prompt = (
            "Please provide a short, clear and concise definition for the marketing term '{}'."
        ).format(topic)
    else:
        if related_links:
            related_links_prompt = (
    """

    In your HTML output, incorporate the following related links into the article text by using relevant anchor text when applicable. 
    If a link is not directly relevant to the text, include it in the 'related terms' section. Here are the related links to incorporate:\n\n{}.

    """
            ).format(", ".join([f"{rl['topic']} ({rl['full path']})" for rl in related_links]))
        else:
            related_links_prompt = ""

        prompt = (
            """

Please write an informative article about the marketing term '{}' following the given outline:\n\n{}\n
Please provide the output in semantic HTML format (there is no need for the H1). {}"""
        ).format(topic, "\n".join(str(sec) for sec in sections), related_links_prompt)

    article = generate_content(api_key, prompt, sections)
    return article


class MyHTMLParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.text = []
        self.current_tag = None
        self.parent_tag = None

    def handle_starttag(self, tag, attrs):
        self.current_tag = tag
        if tag in ["ul", "ol"]:
            self.parent_tag = tag

    def handle_endtag(self, tag):
        if tag == self.parent_tag:
            self.parent_tag = None
        self.current_tag = None

    def handle_data(self, data):
        if self.current_tag in ["h2", "h3", "h4", "p"]:
            self.text.append({"type": self.current_tag, "content": data.strip()})
        elif self.current_tag == "li":
            self.text.append({"type": self.current_tag, "content": data.strip(), "parent": self.parent_tag})

def save_article_as_docx(filename, title, definition, content):
    # Parse the HTML content
    parser = MyHTMLParser()
    parser.feed(content)
    parsed_content = parser.text

    # Create a new DOCX document
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(definition)

    # Add parsed content to the DOCX document
    for item in parsed_content:
        if item["type"] in ["h2", "h3", "h4"]:
            level = int(item["type"][1])
            doc.add_heading(item["content"], level=level)
        elif item["type"] == "p":
            doc.add_paragraph(item["content"])
        elif item["type"] == "li":
            style = "ListBullet" if item["parent"] == "ul" else "ListNumber"
            doc.add_paragraph(item["content"], style=style)

    # Save the document to a file
    doc.save(filename)

def main():
    st.set_page_config(page_title="AI Content Factory", page_icon=None, layout='centered', initial_sidebar_state='expanded')

    st.title("AI Content Factory")
    st.write("A powerful AI-driven content generation tool for creating high-quality articles at scale.")
    
    # Add an expander to provide a detailed overview, benefits, and step-by-step instructions
        
    with st.expander("About this Tool - Content Factory at Your Fingertips"):
        st.markdown("""
Welcome to Content Factory, a powerful content generation tool harnessing the potential of OpenAI's GPT-4 language model. Designed to cater to a diverse array of topics, this tool offers high-quality, informative articles tailored to your specific needs.
""")
    
    with st.expander("Advantages of Content Factory - Unleash the Power of AI"):
        st.markdown("""
The Content Factory boasts numerous benefits for users seeking an efficient, seamless content generation experience:
1. **Tailored content**: By providing specific prompts for each section, the tool generates content that aligns with your unique requirements.
2. **Internal linking**: Automatically generates related links to incorporate within the article, promoting improved SEO.
3. **Scalable**: Effortlessly generates large volumes of content in a timely manner.
4. **User-friendly**: The intuitive interface ensures a smooth, hassle-free experience.

""")
    
    with st.expander("How to Use Content Factory - A Step-by-Step Guide"):
        st.markdown("""
Using the Content Factory is a breeze with this simple step-by-step guide:
1. Run the Content Factory app locally or deploy it on a cloud service.
2. Enter your API key, domain, and upload a CSV file containing your topics and outlines.
3. Adjust the settings as needed, including model, temperature, max tokens, and other options.
4. Click the "Generate Articles" button.
5. Download the generated `.zip` file containing the articles in `.docx` format and a CSV file with the generated content.
6. Review, edit, and utilize the generated content as needed.

""")
    with st.expander("Understanding the Section Start Column Setting"):
        st.markdown("""
The `Section Start Column` setting allows you to specify the column number in your CSV file where the article outline begins. In Python, we start counting from 0, which means that the first column (column A) in the CSV is considered column 0, the second column (column B) is considered column 1, and so on.

For example, if your article outline starts in the 8th column (column H), you should input `7` in the setting, since we begin counting from 0 (column A is 0, column B is 1, ..., column H is 7).

Keep in mind that the last part of your CSV file should be the article outline. This means that no matter which CSV you upload, the outline should be in the final columns of the file, whether it's column 1 or column 50. Ensure that the column immediately following the last section in the article is empty, so the last section of the article should also be the last column in the CSV.

By adjusting the `Section Start Column` setting, you can easily control where the article outline starts in the CSV, making it easier to work with different file structures.
""")
        
    with st.expander("GPT-3.5 Turbo vs. GPT-4 - A Comparison"):
        st.markdown("""
GPT-3.5 Turbo and GPT-4 are both powerful language models, but they differ in terms of speed, cost, and quality of results:

- **Speed**: GPT-3.5 Turbo is faster than GPT-4, generating a 1000-word article in about 50 seconds, while GPT-4 takes around 2 minutes for the same article. This provides users with quicker results when using GPT-3.5 Turbo.

- **Cost**: GPT-4 is more expensive than GPT-3.5 Turbo, and while it may not be as fast, its increased capabilities could justify the higher price for users seeking top-notch content generation.

- **Quality**: GPT-4 delivers better-quality results compared to GPT-3.5 Turbo, which can be an important factor for users who prioritize high-quality content.

""")

    with st.expander("Scaling Content Publishing with Wix - A Comprehensive Approach"):
        st.markdown("""
Scaling content publishing has never been easier, thanks to Wix's robust features and seamless integration with Content Factory. Follow these steps to create a content-rich website:

1. **Create a website with Wix**: Sign up for a Wix account and choose from a variety of stunning templates to build a visually appealing and user-friendly website.

2. **Use Wix's dynamic pages**: Take advantage of Wix's dynamic pages feature to create data-driven, customizable pages that automatically populate with content based on your data.

3. **Upload the CSV file**: Extract the CSV file from the Content Factory-generated `.zip` package and upload it to your Wix site. This file contains all the necessary information, such as topics, definitions, and article content.

4. **Design page layout**: Customize the layout of your dynamic pages to ensure a consistent, visually appealing design that matches your brand identity.

5. **Render the HTML & Schema markup with Velo**: Utilize Velo, Wix's powerful web development platform, to render the HTML and Schema markup generated by Content Factory. This will ensure your content is displayed correctly and optimized for search engines.

By combining the power of Wix and Content Factory, you can streamline the content publishing process and scale your content efforts with ease.

""")

    with st.expander("About Me - Ilay, the Technical SEO Enthusiast"):
        st.markdown("""
Hello! I'm Ilay, a Technical SEO at Wix who loves to cook 🍝 and scale SEO tasks ⏳. Passionate about optimizing content for search engines, I've created Content Factory to help users generate high-quality, AI-generated content effortlessly. I hope you enjoy using this tool as much as I've enjoyed creating it!

""")
        
    with st.sidebar:
        st.header("Settings")

        api_key = st.text_input("API Key:", value="", type="password")
        uploaded_file = st.file_uploader("Upload a CSV file:", type=["csv"])

        domain = st.text_input("URL Path:", value="")

        # Model and settings
        model = st.selectbox("Model:", ["gpt-3.5-turbo", "gpt-4"])
        # top_p = st.slider("Top P:", min_value=0.0, max_value=1.0, value=1.0, step=0.1)
        temperature = st.slider("Temperature:", min_value=0.0, max_value=2.0, value=0.7, step=0.1)
        max_tokens = st.slider("Max Tokens:", min_value=1, max_value=8000, value=2048, step=1)
        presence_penalty = st.slider("Presence Penalty:", min_value=-2.0, max_value=2.0, value=0.2, step=0.1)
        frequency_penalty = st.slider("Frequency Penalty:", min_value=-2.0, max_value=2.0, value=0.2, step=0.1)

        # Add the input field for the section start column
        section_start_col = st.sidebar.number_input("Section Start Column (default is 7)", min_value=1, value=7, step=1)

    if st.button("Generate Articles"):
        if not api_key or not domain or not uploaded_file:
            st.error("Please provide all required inputs (API Key, Domain Name, and CSV File).")
            return

        df = pd.read_csv(uploaded_file)
        df.columns = map(str.lower, df.columns)
        df["url path"] = df["keyword / h1"].apply(create_url_path)
        df["full path"] = df["url path"].apply(lambda x: create_full_path(domain, x))

        topics = df["topic"].tolist()
        h1_keywords = df["keyword / h1"].tolist()
        sections = df.iloc[:, section_start_col-1:].values.tolist()
        
        progress_text = "Generating articles. Please wait..."
        my_bar = st.progress(0, text=progress_text)
        total_items = len(topics) * 2

        definitions = []
        articles = []
        for idx, (topic, sec) in enumerate(zip(topics, sections)):
                related_links = generate_related_links(df, topic)

                definition = generate_article(api_key, topic, sec, related_links, definition_only=True)
                definitions.append(definition)
                time.sleep(7)
                my_bar.progress((((idx + 1) * 2 - 1) / total_items * 100) / 100, text=progress_text)

                article = generate_article(api_key, topic, sec, related_links, definition_only=False)
                articles.append(article)
                time.sleep(7)
                my_bar.progress((((idx + 1) * 2) / total_items * 100) / 100, text=progress_text)

        df["definition"] = definitions
        df["article"] = articles

        os.makedirs("generated_articles", exist_ok=True)

        for idx, (topic, h1_keyword, definition, article) in enumerate(zip(topics, h1_keywords, definitions, articles)):
            docx_filename = f"generated_articles/{topic.replace(' ', '_')}_article.docx"
            save_article_as_docx(docx_filename, h1_keyword, definition, article)

        output_file = "generated_articles/generated_articles.csv"
        df.to_csv(output_file, index=False)

        with zipfile.ZipFile("generated_articles.zip", "w") as zipf:
            for folder, _, filenames in os.walk("generated_articles"):
                for filename in filenames:
                    file_path = os.path.join(folder, filename)
                    zipf.write(file_path, os.path.basename(file_path))

        st.success("Generated articles and definitions added to 'generated_articles.zip'.")

        with open("generated_articles.zip", "rb") as f:
            bytes = f.read()
            b = BytesIO(bytes)
            st.download_button("Download Generated Articles", b, "generated_articles.zip", "application/zip")

if __name__ == "__main__":
    main()
